#https://docs.pinecone.io/docs/semantic-text-search
import docx
import pinecone
import torch
from sentence_transformers import SentenceTransformer
from tqdm.auto import tqdm
from decouple import config

index_name = 'keyword-search'

def insertData(pathToDocs):
    #dataset = load_dataset('quora', split='train')
    #dataset

    #dataset[:5]

    questions = []

    #for record in dataset['questions']:
    #    questions.extend(record['text'])

    doc = docx.Document(pathToDocs)
    for paragraph in doc.paragraphs:
        questions.append(paragraph.text)
    #result = [p.text for p in doc.paragraphs]

    # remove duplicates
    questions = list(set(questions))
    print('\n'.join(questions[:5]))
    print(len(questions))

    device = 'cuda' if torch.cuda.is_available() else 'cpu'
    if device != 'cuda':
        print(f"You are using {device}. This is much slower than using "
              "a CUDA-enabled GPU. If on Colab you can change this by "
              "clicking Runtime > Change runtime type > GPU.")

    model = SentenceTransformer('all-MiniLM-L6-v2', device=device)

    # TODO is this needed?
    #query = 'what strategies are mentioned'

    #xq = model.encode(query)
    #xq.shape

    #_id = '0'
    #metadata = {'text': query}

    #vectors = [(_id, xq, metadata)]

    pinecone.init(
        api_key=config('PINECONE_API_KEY'),  # find at app.pinecone.io
        environment=config('PINECONE_ENVIRONMENT') # next to api key in console
    )

    if index_name not in pinecone.list_indexes():
        pinecone.create_index(
            name=index_name,
            dimension=model.get_sentence_embedding_dimension(),
            metric='cosine'
        )

    # now connect to the index
    index = pinecone.GRPCIndex(index_name)

    batch_size = 128

    for i in tqdm(range(0, len(questions), batch_size)):
        # find end of batch
        i_end = min(i+batch_size, len(questions))
        # create IDs batch
        ids = [str(x) for x in range(i, i_end)]
        # create metadata batch
        metadatas = [{'text': text} for text in questions[i:i_end]]
        # create embeddings
        xc = model.encode(questions[i:i_end])
        # create records list for upsert
        records = zip(ids, xc, metadatas)
        # upsert to Pinecone
        index.upsert(vectors=records)

    # check number of records in the index
    index.describe_index_stats()

def deletIndex():
    pinecone.delete_index(index_name)

if __name__ == '__main__':
    pathToWordFile = "C:\\Users"
    insertData(pathToWordFile + "\\Strategic_Alignment_Process.docx")